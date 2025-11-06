import streamlit as st
from llama_index.core import VectorStoreIndex, Settings, Document
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.openai import OpenAI
from llama_index.core.memory import ChatMemoryBuffer
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import pandas as pd
import time
import os

# ---------------------------------------------------------
# ✅ Page Setup + Load CSS
# ---------------------------------------------------------
st.set_page_config(page_title="AI Document Chatbot", layout="wide", page_icon="🤖")

def local_css(file_name):
    with open(file_name, "r") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

local_css("style.css")     

st.title("🧠 AI Document Chatbot")
st.caption("Powered by LlamaIndex + OpenRouter + HuggingFace Embeddings")

# ---------------------------------------------------------
# ✅ Load .env (OpenRouter API KEY)
# ---------------------------------------------------------
from dotenv import load_dotenv
load_dotenv()

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")

if not OPENROUTER_API_KEY:
    st.error("❌ OPENROUTER_API_KEY not found in .env. Please add it.")
else:
    os.environ["OPENAI_API_KEY"] = OPENROUTER_API_KEY
    os.environ["OPENAI_API_BASE"] = "https://openrouter.ai/api/v1"
    os.environ["OPENAI_API_MODEL"] = "gpt-3.5-turbo"


# ---------------------------------------------------------
# Step 1: File Upload
# ---------------------------------------------------------
uploaded_files = st.file_uploader(
    "📂 Upload your files (.txt, .pdf, .docx, .csv)",
    type=["txt", "pdf", "docx", "csv"],
    accept_multiple_files=True
)

docs = []

if uploaded_files:
    with st.spinner("🔍 Processing uploaded files..."):
        for uploaded_file in uploaded_files:
            file_name = uploaded_file.name

            # ---- TXT ----
            if file_name.endswith(".txt"):
                text = uploaded_file.read().decode("utf-8")
                docs.append(Document(text=text, metadata={"file": file_name}))

            # ---- PDF ----
            elif file_name.endswith(".pdf"):
                reader = PdfReader(uploaded_file)
                text = "\n".join([
                    page.extract_text()
                    for page in reader.pages
                    if page.extract_text()
                ])
                docs.append(Document(text=text, metadata={"file": file_name}))

            # ---- DOCX ----
            elif file_name.endswith(".docx"):
                doc = DocxDocument(uploaded_file)
                text = "\n".join([p.text for p in doc.paragraphs])
                docs.append(Document(text=text, metadata={"file": file_name}))

            # ---- CSV ----
            elif file_name.endswith(".csv"):
                df = pd.read_csv(uploaded_file)
                text = df.to_string(index=False)
                docs.append(Document(text=text, metadata={"file": file_name}))

        time.sleep(1.5)

    success = st.success(f"✅ Loaded {len(docs)} documents successfully!")
    time.sleep(2)
    success.empty()

# ---------------------------------------------------------
# Step 2: Embeddings + Chat Engine
# ---------------------------------------------------------
if docs:
    with st.spinner("⚙️ Building vector index..."):
        Settings.embed_model = HuggingFaceEmbedding(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        Settings.llm = OpenAI(
            model="gpt-3.5-turbo",
            temperature=0
        )

        index = VectorStoreIndex.from_documents(docs)
        memory = ChatMemoryBuffer.from_defaults(token_limit=3900)
        chat_engine = index.as_chat_engine(
            chat_mode="context",
            memory=memory,
            llm=Settings.llm
        )
        time.sleep(1.5)

    ready = st.success("🚀 Chat engine ready! Start chatting below 👇")
    time.sleep(2)
    ready.empty()

    # ---------------------------------------------------------
    # Step 3: Chat Interface
    # ---------------------------------------------------------
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    user_input = st.chat_input("💬 Ask a question about your documents...")

    if user_input:
        with st.chat_message("user"):
            st.markdown(user_input)

        response = chat_engine.chat(user_input)

        with st.chat_message("assistant"):
            st.markdown(response.response)

        st.session_state.chat_history.append((user_input, response.response))

    # ---------------------------------------------------------
    # Step 4: Chat History (Sidebar)
    # ---------------------------------------------------------
    if st.session_state.chat_history:
        st.sidebar.header("🕒 Chat History")

        if st.sidebar.button("🗑 Clear Chat"):
            st.session_state.chat_history = []
            st.rerun()

        for i, (q, a) in enumerate(reversed(st.session_state.chat_history), 1):
            st.sidebar.markdown(f"**Q{i}:** {q}")
            st.sidebar.markdown(f"**A{i}:** {a}")

else:
    st.warning("⬆️ Please upload at least one document to begin chatting.")
